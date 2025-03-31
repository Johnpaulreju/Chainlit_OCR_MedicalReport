

# import chainlit as cl
# import pdfplumber
# import docx
# import openai
# import os
# import re
# from dotenv import load_dotenv

# # Load API key from .env file
# load_dotenv()
# openai.api_key = os.getenv("OPENAI_API_KEY")

# MAX_TOKENS = 6000  # Keep a buffer to avoid exceeding GPT's limit

# @cl.on_message
# async def process_message(message: cl.Message):
#     files = None

#     while files is None:
#         files = await cl.AskFileMessage(
#             content="Please upload a document (PDF, DOCX, or TXT).",
#             accept=["application/pdf", "application/vnd.openxmlformats-officedocument.wordprocessingml.document", "text/plain"]
#         ).send()

#     file = files[0]  
#     file_path = file.path  
#     extracted_text = extract_text(file_path, file.name)

#     # Clean and filter extracted text
#     filtered_text = filter_unnecessary_info(extracted_text)

#     # Ensure text is within GPT's token limit
#     if len(filtered_text.split()) > MAX_TOKENS:
#         filtered_text = summarize_text(filtered_text)  # Summarize if too long

#     structured_response = analyze_with_llm(filtered_text)

#     await cl.Message(content=structured_response).send()

# def extract_text(file_path, file_name):
#     """Extracts text from PDFs, DOCX, and TXT files."""
#     text = ""

#     if file_name.endswith(".pdf"):
#         with pdfplumber.open(file_path) as pdf:
#             for page in pdf.pages:
#                 page_text = page.extract_text()
#                 if page_text:
#                     text += clean_text(page_text) + "\n\n"

#                 # Extract tables
#                 tables = page.extract_tables()
#                 for table in tables:
#                     text += format_table(table) + "\n\n"

#     elif file_name.endswith(".docx"):
#         doc = docx.Document(file_path)
#         text = "\n".join([para.text for para in doc.paragraphs])

#     elif file_name.endswith(".txt"):
#         with open(file_path, "r", encoding="utf-8") as f:
#             text = f.read()

#     else:
#         text = "Unsupported file format."

#     return text

# def clean_text(text):
#     """Cleans extracted text by removing extra spaces and fixing line breaks."""
#     lines = text.split("\n")
#     cleaned_lines = [line.strip() for line in lines if line.strip()]
#     return "\n".join(cleaned_lines)

# def format_table(table):
#     """Formats tables into structured text."""
#     formatted_table = "\n".join([" | ".join(row) for row in table if any(row)])
#     return f"\n--- Table Extracted ---\n{formatted_table}\n--- End of Table ---\n"

# def filter_unnecessary_info(text):
#     """Removes irrelevant details, keeping only Patient Name & Test Results."""
#     lines = text.split("\n")

#     # Keywords to filter out
#     unwanted_patterns = [
#         r"\bScan QR code\b", r"\breport authenticity\b", r"\bPassport No\b", 
#         r"\bLABORATORY TEST REPORT\b", r"\bClient Name\b", r"\bLab Id\b",
#         r"\bRef. Id\b", r"\bPrinted On\b", r"\bApproved on\b", r"\bProcess At\b",
#         r"\bSample Information\b", r"\bLocation\b", r"\bRef. By\b"
#     ]

#     # Keep only Patient Name and Test Results
#     cleaned_lines = []
#     keep_section = False
#     for line in lines:
#         # Remove unwanted lines
#         if any(re.search(pattern, line, re.IGNORECASE) for pattern in unwanted_patterns):
#             continue

#         # If "Test Results" or similar section starts, keep content
#         if "Test Result" in line or "Blood Count" in line or "Lipid Profile" in line:
#             keep_section = True

#         if keep_section:
#             cleaned_lines.append(line)

#     return "\n".join(cleaned_lines)

# def summarize_text(text):
#     """Summarizes large text before sending it to GPT to fit within token limits."""
#     client = openai.OpenAI()

#     summary_prompt = f"""
#     The following text is too long to process. Please summarize it while keeping key details:

#     {text[:8000]}  # Taking only the first 8000 characters for context.

#     Make sure the summary is under {MAX_TOKENS} tokens.
#     """

#     response = client.chat.completions.create(
#         model="gpt-4",
#         messages=[
#             {"role": "system", "content": "You are an expert at summarizing long documents."},
#             {"role": "user", "content": summary_prompt}
#         ]
#     )

#     return response.choices[0].message.content

# def analyze_with_llm(text):
#     """Sends extracted text to GPT for structured analysis using the latest OpenAI API."""
#     client = openai.OpenAI()

#     prompt = f"""
#     The following text is extracted from a medical report. Please analyze and structure it properly.

#     1. **Identify key sections** (Patient Name, Test Results).
#     2. **Format extracted tables properly**.
#     3. **Summarize the important details in a structured format**.

#     Here is the extracted text:
#     {text}

#     Please format this properly into an easy-to-read report.
#     """

#     response = client.chat.completions.create(
#         model="gpt-4",
#         messages=[
#             {"role": "system", "content": "You are an AI that formats and analyzes medical reports."},
#             {"role": "user", "content": prompt}
#         ]
#     )

#     return response.choices[0].message.content


import chainlit as cl
import pdfplumber
import docx
import openai
import os
import re
from dotenv import load_dotenv

# Load API key from .env file
load_dotenv()
openai.api_key = os.getenv("OPENAI_API_KEY")

MAX_TOKENS = 6000  # Buffer to avoid exceeding GPT's limit

@cl.on_message
async def process_message(message: cl.Message):
    """Handles user message and extracts text from uploaded files."""
    
     # ✅ Check if user uploaded a file via spontaneous upload
    uploaded_files = [el for el in message.elements if el.type == "file"]

    if not uploaded_files:
        await cl.Message(content="Please upload a medical report (PDF, DOCX, or TXT).").send()
        return
    
    file = uploaded_files[0]  # Get the first uploaded file
    file_path = file.path  
    extracted_text = extract_text(file_path, file.name)
    print("extracting completed")
    print("\n")
    # Clean and filter extracted text dynamically
    filtered_text = filter_repetitions(extracted_text)
    print("filteration completed")
    print("\n")

    # Ensure text is within GPT's token limit
    structured_text = handle_large_text(filtered_text)
    print("structuring completed")
    print("\n")

    # ✅ Fix: **Await GPT response properly**
    structured_response = analyze_with_llm(structured_text)  # ✅ Remove 'await'
    print("analyzing completed",structured_response)
    print("\n")

    print("📤 Sending response to Chainlit UI...")
    if structured_response:
        print("📤 Sending response to Chainlit UI...")
        await cl.Message(structured_response).send()
        print("✅ Response sent successfully!")
    else:
        print("⚠️ Empty response received from GPT. Not sending.")


def extract_text(file_path, file_name):
    """Extracts text from PDFs, DOCX, and TXT files while preserving table data."""
    text = ""

    if file_name.endswith(".pdf"):
        with pdfplumber.open(file_path) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text += clean_text(page_text) + "\n\n"

                # Extract tables properly
                tables = page.extract_tables()
                for table in tables:
                    text += format_table(table) + "\n\n"

    elif file_name.endswith(".docx"):
        doc = docx.Document(file_path)
        text = "\n".join([para.text for para in doc.paragraphs])

    elif file_name.endswith(".txt"):
        with open(file_path, "r", encoding="utf-8") as f:
            text = f.read()

    else:
        text = "Unsupported file format."

    return text

def clean_text(text):
    """Cleans extracted text by removing extra spaces and fixing line breaks."""
    lines = text.split("\n")
    cleaned_lines = [line.strip() for line in lines if line.strip()]
    return "\n".join(cleaned_lines)

def format_table(table):
    """Ensures extracted tables retain structure while handling NoneType values."""
    formatted_table = "\n".join([" | ".join([str(cell) if cell is not None else "" for cell in row]) for row in table if any(row)])
    return f"\n--- Table Extracted ---\n{formatted_table}\n--- End of Table ---\n"


def filter_repetitions(text):
    """Removes duplicate sections dynamically while keeping all unique test results."""
    lines = text.split("\n")

    # Track seen sections and subtests to prevent full repetition but retain new test results
    seen_sections = {}
    filtered_lines = []
    current_section = None
    section_content = []

    for line in lines:
        # Detect medical test sections dynamically
        section_match = re.search(r"^[IVXLCDM]+\.?\s?(.*)", line)  # Matches Roman numeral test sections
        if section_match:
            # If we were processing a previous section, check if it has unique data
            if current_section:
                section_data = "\n".join(section_content).strip()
                if current_section not in seen_sections or seen_sections[current_section] != section_data:
                    filtered_lines.append(current_section)  # Add section header
                    filtered_lines.extend(section_content)  # Add its unique data
                    seen_sections[current_section] = section_data  # Save it
                
            # Start a new section
            current_section = line.strip()
            section_content = []
            continue  # Skip adding section header immediately, wait for its data

        # If inside a section, add content to it
        if current_section:
            section_content.append(line)
        else:
            filtered_lines.append(line)  # If not inside a section, add normally

    # Ensure the last section is added
    if current_section:
        section_data = "\n".join(section_content).strip()
        if current_section not in seen_sections or seen_sections[current_section] != section_data:
            filtered_lines.append(current_section)  # Add section header
            filtered_lines.extend(section_content)  # Add its unique data

    return "\n".join(filtered_lines)

def handle_large_text(text):
    """Handles text exceeding GPT's limit by summarizing while keeping medical details."""
    if len(text.split()) > MAX_TOKENS:
        return summarize_text(text)
    return text

def summarize_text(text):
    """Summarizes large text before sending it to GPT to fit within token limits."""
    client = openai.OpenAI()

    summary_prompt = f"""
    The following medical report is too long. Summarize it while keeping key test details:

    {text[:8000]}

    Ensure the summary is under {MAX_TOKENS} tokens.
    """

    response = client.chat.completions.create(
        model="gpt-4o",
        messages=[
            {"role": "system", "content": "You are an expert at summarizing medical reports while preserving test details."},
            {"role": "user", "content": summary_prompt}
        ]
    )

    return response.choices[0].message.content

def analyze_with_llm(text):
    """Formats extracted text into a structured Markdown table for proper Chainlit rendering."""
    client = openai.OpenAI()

    prompt = f"""
    The following text is extracted from a medical report. Format it into a **clean Markdown table**.

    1. **Ensure ALL test sections and subtests are included.**
    2. **Check for missing tests (e.g., 'Microalbumin', 'Total WBC and Differential Count').**
    3. **If a test section appears multiple times, include only new test values, not the entire section again.**
    4. **Extract relevant test names, results, units, and reference intervals.**
    5. **Highlight abnormal values using `H` (High) or `L` (Low).**
    6. **Ensure output follows this Markdown table format:**  

    | **Test**                     | **Result**    | **Unit**          | **Biological Ref. Interval** |
    |------------------------------|--------------|-------------------|------------------------------|
    | Hemoglobin                   | 14.5         | g/dL              | 13.0 - 16.5                  |
    | RBC Count                    | 4.79         | million/cmm       | 4.5 - 5.5                    |

    7. **Return the table as clean Markdown (NO triple backticks ` ``` `).**
    8. **Summarize findings after the table.**

    Extracted Medical Report Data:
    {text}
    """

    response = client.chat.completions.create(
        model="gpt-4o",
        messages=[
            {"role": "system", "content": "You are an AI that extracts and formats medical test results into structured tables using Markdown."},
            {"role": "user", "content": prompt}
        ]
    )

    return response.choices[0].message.content



def analyze_with_llm(text):
    """Formats extracted text into two structured Markdown tables: 
    1. Full Report (All test results)
    2. Issues Found (Abnormal or Borderline Values)"""
    
    client = openai.OpenAI()

    prompt = f"""
    The following text is extracted from a medical report. Your task is to:
    
    1. **Create a table with ALL test results**.
    2. **Create a second table called "Issues Found"**:
       - Include only values that are **above or below reference limits**.
       - Also include **borderline values** (close to limit but within range).
       - Mark abnormalities with:
         - 🔴 **High (H) or Low (L)**
         - 🔶 **Borderline**
    3. **Use this Markdown table format** (NO code blocks! Just plain Markdown):
    
    ### Full Medical Report:
    | **Test**                     | **Result**    | **Unit**          | **Biological Ref. Interval** |
    |------------------------------|--------------|-------------------|------------------------------|
    | Hemoglobin                   | 14.5         | g/dL              | 13.0 - 16.5                  |
    | RBC Count                    | 4.79         | million/cmm       | 4.5 - 5.5                    |
    
    ### Issues Found:
    | **Test**                     | **Result**    | **Unit**          | **Biological Ref. Interval** |
    |------------------------------|--------------|-------------------|------------------------------|
    | Triglyceride                 | **168.0** 🔴 | mg/dL             | **Normal <150**              |
    | HbA1c                        | **7.10** 🔴  | %                 | **Diabetes >6.5%**           |

    **Extracted Medical Report Data:**
    {text}
    """

    response = client.chat.completions.create(
        model="gpt-4o",
        messages=[
            {"role": "system", "content": "You are an AI that extracts and formats medical test results into structured tables using Markdown."},
            {"role": "user", "content": prompt}
        ]
    )

    return response.choices[0].message.content

